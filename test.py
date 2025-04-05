import os
import datetime
import docx
from langchain_ollama import OllamaLLM
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_unstructured import UnstructuredLoader
from langchain.chains import RetrievalQA
from playsound import playsound
from TTS.api import TTS
import speech_recognition as sr
import threading
import time
import re

os.environ["FAISS_NO_AVX2"] = "1"

transcript_buffer = []
context_chunks = []
CHUNK_SIZE = 20
TRANSCRIPT_FOLDER = "transcripts"
os.makedirs(TRANSCRIPT_FOLDER, exist_ok=True)

ollama = OllamaLLM(base_url='http://localhost:11434', model="llama3")
tts = TTS(model_name="tts_models/en/ljspeech/glow-tts")
r = sr.Recognizer()
lock = threading.Lock()
history_vectorstores = []
summarized_once = False


def correct_text(text):
    return re.sub(r"\bim\b", "I'm", text, flags=re.IGNORECASE)


def speak(text):
    try:
        if not text.strip():
            print("Nothing to speak.")
            return

        output_file = "output.wav"
        if os.path.exists(output_file):
            try:
                os.remove(output_file)
            except Exception as e:
                print(f"Couldn't delete old audio: {e}")
                return

        tts.tts_to_file(text=text, file_path=output_file)

        for _ in range(5):
            try:
                playsound(output_file)
                os.remove(output_file)
                break
            except PermissionError:
                print("Waiting for file to be released...")
                time.sleep(0.5)
        else:
            print("Could not play audio. Skipping.")

    except Exception as e:
        print(f"TTS Error: {e}")


def listen():
    with sr.Microphone() as source:
        print("Listening...")
        audio = r.listen(source)
    try:
        text = r.recognize_google(audio)
        print(f"You said: {text}")
        return text
    except sr.UnknownValueError:
        print("Could not understand audio")
        return None
    except sr.RequestError as e:
        print(f"Google Speech Recognition error: {e}")
        return None


def get_today_file_path():
    date_str = datetime.datetime.now().strftime("%Y-%m-%d")
    return os.path.join(TRANSCRIPT_FOLDER, f"Class_{date_str}.docx")


def save_to_docx(text):
    file_path = get_today_file_path()
    if not os.path.exists(file_path):
        doc = docx.Document()
        doc.add_heading('Class Notes', level=1)
        doc.save(file_path)

    doc = docx.Document(file_path)
    doc.add_paragraph(text)
    doc.save(file_path)


def convert_doc_to_vector_db(file_path):
    print("Converting .docx to vector DB...")
    loader = UnstructuredLoader(file_path)
    documents = loader.load()

    documents = [doc for doc in documents if hasattr(doc, 'page_content') and doc.page_content.strip() != ""]
    if not documents:
        print("No valid content found in document.")
        speak("Unable to extract any valid content from the document.")
        return None

    for i, doc in enumerate(documents):
        print(f"Document {i+1}: {doc.page_content[:100]}...")

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    texts = [doc.page_content for doc in documents]
    metadatas = [doc.metadata for doc in documents]

    import faiss
    import numpy as np
    vectors = embeddings.embed_documents(texts)
    dimension = len(vectors[0])
    index = faiss.IndexFlatL2(dimension)
    index.add(np.array(vectors).astype("float32"))

    from langchain_community.vectorstores.faiss import InMemoryDocstore
    docstore = InMemoryDocstore({i: documents[i] for i in range(len(documents))})
    vectorstore = FAISS(embedding_function=embeddings, index=index, index_to_docstore_id=dict(enumerate(range(len(texts)))), docstore=docstore)
    return vectorstore


def summarize_today_class(vectorstore):
    global summarized_once
    if summarized_once:
        speak("I have already given the summary. Say 'repeat summary' if you want to hear it again.")
        return

    print("Generating summary using LLM...")
    qa = RetrievalQA.from_chain_type(llm=ollama, chain_type="stuff", retriever=vectorstore.as_retriever())
    try:
        response = qa.invoke({"query": "Give me a detailed summary of all the class content."})
    except Exception as e:
        print(f"Summary generation failed: {e}")
        speak("I was unable to generate the summary due to an error.")
        return

    summary = response.get("result", "").strip() if isinstance(response, dict) else str(response).strip()

    if not summary:
        print("Summary is empty. Skipping.")
        speak("Sorry, I could not generate a summary.")
        return

    print("Summary generated:")
    print(summary)
    speak(summary)
    summarized_once = True


def handle_questions(today_vectorstore):
    combined_docs = []
    for vs in [today_vectorstore] + history_vectorstores:
        combined_docs.extend(vs.similarity_search(""))

    combined_docs = [doc for doc in combined_docs if hasattr(doc, 'page_content') and doc.page_content.strip()]
    if not combined_docs:
        speak("There is no content available for Q and A.")
        return

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    merged_vectorstore = FAISS.from_documents(combined_docs, embeddings)
    qa = RetrievalQA.from_chain_type(llm=ollama, chain_type="stuff", retriever=merged_vectorstore.as_retriever())

    speak("You may now ask questions. I will wait for 2 minutes.")
    start_time = time.time()

    while True:
        question = listen()
        if question:
            if question.lower() in ["no", "exit", "quit"]:
                speak("Ending Q&A session.")
                break
            elif "repeat summary" in question.lower():
                summarize_today_class(today_vectorstore)
                continue

            answer = qa.invoke({"query": question})
            answer_text = answer.get("result", "")
            if answer_text.strip():
                print(answer_text)
                speak(answer_text)
            else:
                speak("Sorry, I don't know the answer.")
            start_time = time.time()
        else:
            if time.time() - start_time > 120:
                speak("No questions received. Ending the class now.")
                break


def main():
    print("Class started. Listening...")
    today_file_path = get_today_file_path()
    today_vectorstore = None

    trigger_phrases = [
        "end of class", "class over", "give me the summary",
        "summarize today", "class is finished", "class is end",
        "give me the summarization audio", "summarize today's class",
        "class ended", "finish the class"
    ]

    while True:
        text = listen()
        if text:
            if any(phrase in text.lower() for phrase in trigger_phrases):
                print("Trigger phrase detected. Ending class session.")
                break

            cleaned_text = correct_text(text)
            with lock:
                transcript_buffer.append(cleaned_text)
                save_to_docx(cleaned_text)

                all_text = " ".join(transcript_buffer)
                if len(all_text.split()) >= CHUNK_SIZE:
                    print("Enough tokens collected. Creating/updating vector DB...")
                    today_vectorstore = convert_doc_to_vector_db(today_file_path)
                    transcript_buffer.clear()

    if today_vectorstore is None:
        print("Generating fallback vector DB for small transcript...")
        today_vectorstore = convert_doc_to_vector_db(today_file_path)

    if today_vectorstore:
        history_vectorstores.append(today_vectorstore)
        summarize_today_class(today_vectorstore)
        handle_questions(today_vectorstore)
    else:
        speak("I couldn't understand the class content well enough to summarize.")


if __name__ == "__main__":
    main()
